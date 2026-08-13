import os
import sys
import docx
from src.document_processor import DocumentProcessor
from src.pii_detector import PIIDetector

def extract_all_xml_text(doc_path):
    """Extract all raw text and relationship targets from all XML parts in the DOCX package."""
    import zipfile
    text_parts = []
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        for name in zip_ref.namelist():
            if name.endswith('.xml') or name.endswith('.rels'):
                try:
                    content = zip_ref.read(name)
                    text_parts.append(content.decode('utf-8'))
                except Exception:
                    pass
    return "\n".join(text_parts)

def extract_all_text(doc_path):
    """Utility to extract all text from a docx file (paragraphs, tables, headers, footers)."""
    doc = docx.Document(doc_path)
    text_parts = []
    
    # Extract headers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            for t in section.header.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                text_parts.append(p.text)
                                
    # Extract main paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text)
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text_parts.append(p.text)
                        
    # Extract footers
    for section in doc.sections:
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            for t in section.footer.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                text_parts.append(p.text)
                                
    return "\n".join(text_parts)

def main(input_file, output_file):
    print("=" * 60)
    print("Scaler AI Labs - PII Redaction Pipeline starting...")
    print(f"Input file: {input_file}")
    print(f"Output file: {output_file}")
    print("=" * 60)
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' does not exist.")
        sys.exit(1)
        
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # 1. Run the Document Redaction
    processor = DocumentProcessor(input_file, output_file)
    stats = processor.process()
    
    print("\n--- Redaction Finished Successfully ---")
    print(f"Paragraphs processed: {stats['paragraphs_processed']}")
    print(f"Tables processed: {stats['tables_processed']}")
    print(f"Table cells processed: {stats['cells_processed']}")
    print(f"Headers processed: {stats['headers_processed']}")
    print(f"Footers processed: {stats['footers_processed']}")
    print(f"Total replacements made: {stats['replacements']}")
    print("PII Detections by Category:")
    for cat, count in sorted(stats['detections_by_category'].items()):
        print(f"  - {cat}: {count}")
        
    # 2. Post-Redaction Verification Check
    print("\n--- Running Post-Redaction Verification ---")
    redacted_text = extract_all_text(output_file)
    package_text = extract_all_xml_text(output_file)
    
    mapping = processor.redactor.mapping
    redacted_text_lower = redacted_text.lower()
    package_text_lower = package_text.lower()
    
    leaks_visible = []
    leaks_package = []
    
    for original_pii, replacement in mapping.items():
        if len(original_pii) < 4:
            continue
        # Check visible text
        if original_pii.lower() in redacted_text_lower:
            leaks_visible.append((original_pii, replacement))
        # Check complete XML / ZIP package
        if original_pii.lower() in package_text_lower:
            leaks_package.append((original_pii, replacement))
            
    # Check for original mailto: targets in the XML package
    import re
    mailto_matches = re.findall(r'mailto:([a-zA-Z0-9\._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', package_text, re.IGNORECASE)
    original_emails = {k.lower() for k, v in mapping.items() if '@' in k}
    leaks_mailto = 0
    for email in mailto_matches:
        if email.lower() in original_emails:
            leaks_mailto += 1
            
    # Run PII Detector again on the redacted text to spot uncaught PII
    re_detector = PIIDetector()
    remaining_pii = re_detector.detect_all(redacted_text)
    
    # Filter remaining PII to only include high confidence ones that are not already our fake replacements
    actual_uncaught = []
    fake_vals = set(mapping.values())
    for p in remaining_pii:
        val = p["original"]
        # Skip if it is a substring of any synthetic replacement
        if any(val in fv for fv in fake_vals):
            continue
        # Skip if it's too short
        if len(val) < 4:
            continue
        actual_uncaught.append(p)
        
    print(f"Verification Results:")
    print(f"  - Original PII in visible text: {len(leaks_visible)}")
    for val, repl in leaks_visible[:10]:
        print(f"    WARNING: Original PII '{val}' is still present in text! (Should be '{repl}')")
        
    print(f"  - Original PII in DOCX XML/package: {len(leaks_package)}")
    for val, repl in leaks_package[:10]:
        print(f"    WARNING: Original PII '{val}' is still present in XML package! (Should be '{repl}')")
        
    print(f"  - Original mailto targets: {leaks_mailto}")
    
    print(f"  - Remaining high-confidence uncaught PII: {len(actual_uncaught)}")
    for p in actual_uncaught[:10]:
        print(f"    WARNING: Potential uncaught {p['type']}: '{p['original']}'")
        
    # Save a verification report summary in output folder
    verification_summary_path = os.path.join(os.path.dirname(output_file), "verification_summary.txt")
    with open(verification_summary_path, "w", encoding="utf-8") as f:
        f.write("=== POST-REDACTION VERIFICATION SUMMARY ===\n")
        f.write(f"Total original PII replaced: {len(mapping)}\n")
        f.write(f"Original PII in visible text: {len(leaks_visible)}\n")
        f.write(f"Original PII in DOCX XML/package: {len(leaks_package)}\n")
        f.write(f"Original mailto targets: {leaks_mailto}\n")
        f.write(f"Potential uncaught PII remaining: {len(actual_uncaught)}\n\n")
        if leaks_package:
            f.write("--- LEAKED ENTITIES IN PACKAGE ---\n")
            for val, repl in leaks_package:
                f.write(f"Original: {val} | Replacement: {repl}\n")
        if actual_uncaught:
            f.write("--- POTENTIAL UNCAUGHT PII ---\n")
            for p in actual_uncaught:
                f.write(f"Type: {p['type']} | Value: {p['original']}\n")
                
    print(f"\nVerification report saved to: {verification_summary_path}")
    print("=" * 60)
    
    return {
        "stats": stats,
        "mapping_count": len(mapping),
        "leaks_count": len(leaks_package),
        "uncaught_count": len(actual_uncaught)
    }

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="PII Redactor Main CLI")
    parser.add_argument("--input", default="pii-redaction-tool/input/Red Herring Prospectus.docx", help="Input file path")
    parser.add_argument("--output", default="pii-redaction-tool/output/Redacted_Red_Herring_Prospectus.docx", help="Output file path")
    args = parser.parse_args()
    main(args.input, args.output)
