import docx
import re
from src.pii_detector import PIIDetector
from src.redactor import PIIRedactor

def replace_span_in_runs(runs, start, end, replacement):
    """
    Replaces characters from start to end index in a list of runs,
    preserving formatting.
    """
    offset = 0
    start_run_idx = -1
    start_run_offset = -1
    end_run_idx = -1
    end_run_offset = -1
    
    for i, run in enumerate(runs):
        run_len = len(run.text)
        if start_run_idx == -1 and offset <= start < offset + run_len:
            start_run_idx = i
            start_run_offset = start - offset
        if end_run_idx == -1 and offset < end <= offset + run_len:
            end_run_idx = i
            end_run_offset = end - offset
        offset += run_len
        
    # Handle end of text boundary
    if start == offset and start_run_idx == -1 and len(runs) > 0:
        start_run_idx = len(runs) - 1
        start_run_offset = len(runs[-1].text)
    if end == offset and end_run_idx == -1 and len(runs) > 0:
        end_run_idx = len(runs) - 1
        end_run_offset = len(runs[-1].text)
        
    if start_run_idx == -1 or end_run_idx == -1:
        return False
        
    if start_run_idx == end_run_idx:
        run = runs[start_run_idx]
        run.text = run.text[:start_run_offset] + replacement + run.text[end_run_offset:]
    else:
        start_run = runs[start_run_idx]
        start_run.text = start_run.text[:start_run_offset] + replacement
        for i in range(start_run_idx + 1, end_run_idx):
            runs[i].text = ""
        end_run = runs[end_run_idx]
        end_run.text = end_run.text[end_run_offset:]
    return True

class DocumentProcessor:
    def __init__(self, input_path, output_path):
        self.input_path = input_path
        self.output_path = output_path
        self.detector = PIIDetector()
        self.redactor = PIIRedactor()
        
        # Summary statistics
        self.stats = {
            "paragraphs_processed": 0,
            "tables_processed": 0,
            "cells_processed": 0,
            "headers_processed": 0,
            "footers_processed": 0,
            "replacements": 0,
            "detections_by_category": {}
        }

    def process(self):
        """Load, process, redact, and save the DOCX document."""
        print(f"Loading document: {self.input_path}")
        doc = docx.Document(self.input_path)
        
        # 1. Process Main Stories (Paragraphs)
        self.process_paragraphs(doc.paragraphs)
        
        # 2. Process Tables
        self.process_tables(doc.tables)
        
        # 3. Process Headers & Footers in all sections (deduplicated to prevent linked double-processing)
        processed_headers = set()
        processed_footers = set()
        
        for section_idx, section in enumerate(doc.sections):
            if section.header and section.header not in processed_headers:
                processed_headers.add(section.header)
                self.stats["headers_processed"] += 1
                self.process_paragraphs(section.header.paragraphs)
                self.process_tables(section.header.tables)
            if section.footer and section.footer not in processed_footers:
                processed_footers.add(section.footer)
                self.stats["footers_processed"] += 1
                self.process_paragraphs(section.footer.paragraphs)
                self.process_tables(section.footer.tables)
                
        print(f"Saving redacted document to: {self.output_path}")
        doc.save(self.output_path)
        return self.stats

    def process_paragraphs(self, paragraphs):
        """Redact a list of paragraphs using a single-replacement loop to prevent offset shifts."""
        for p in paragraphs:
            self.stats["paragraphs_processed"] += 1
            
            loop_count = 0
            while loop_count < 100:
                loop_count += 1
                text = p.text
                if not text.strip():
                    break
                    
                matches = self.detector.detect_all(text)
                unredacted = []
                for m in matches:
                    norm_val = re.sub(r"\s+", " ", m["original"].strip())
                    # Bidirectional overlap check: skip if fake contains match or match contains fake
                    if any(fv in norm_val or norm_val in fv for fv in self.redactor.generated_values):
                        continue
                    unredacted.append(m)
                    
                if not unredacted:
                    break
                    
                # Process the first unredacted match
                m = unredacted[0]
                repl = self.redactor.get_replacement(m["original"], m["type"])
                success = replace_span_in_runs(p.runs, m["start"], m["end"], repl)
                if success:
                    self.stats["replacements"] += 1
                    self.stats["detections_by_category"][m["type"]] = \
                        self.stats["detections_by_category"].get(m["type"], 0) + 1
                else:
                    # If run replacement failed, break to prevent infinite loop
                    break

    def process_tables(self, tables):
        """Redact a list of tables, utilizing column-header context heuristics."""
        from src.pii_detector import BLACKLISTED_NAME_TERMS
        
        def validate_cell_text(text, pii_type):
            raw_text = text.strip()
            if not raw_text or raw_text in ["N.A.", "NA", "[●]", "●"]:
                return False
                
            # Clean footnote markers and special symbols for validation checks
            clean_text = re.sub(r"[\*\#\^\&\@\(\)\[\]●\-\–\—\xa0]+", " ", raw_text)
            clean_text = re.sub(r"\s+", " ", clean_text).strip()
            
            if pii_type == "EMAIL":
                return "@" in clean_text and "." in clean_text
                
            elif pii_type == "PHONE":
                digits = re.sub(r"\D", "", clean_text)
                if len(digits) < 8 or len(digits) > 13:
                    return False
                letters = re.sub(r"[^a-zA-Z]", "", clean_text)
                if len(letters) > 3:
                    return False
                return True
                
            elif pii_type == "ADDRESS":
                if len(clean_text) < 15:
                    return False
                indicators = ["pune", "mumbai", "bhopal", "maharashtra", "delhi", "road", "marg", "street", "plot", "gat", "flat", "building", "floor", "office", "village", "taluka", "district", "india", "pincode", "pin –", "pin -"]
                text_lower = clean_text.lower()
                if not any(ind in text_lower for ind in indicators) and not re.search(r"\b\d{3}\s?\d{3}\b", clean_text):
                    return False
                return True
                
            elif pii_type == "FULL_NAME":
                if re.search(r"\d", clean_text):
                    return False
                if clean_text in BLACKLISTED_NAME_TERMS or any(term in clean_text for term in BLACKLISTED_NAME_TERMS):
                    return False
                words = clean_text.split()
                # Check capitalization on cleaned words
                if not words or not all(w[0].isupper() for w in words if w.lower() not in ["and", "of", "&"]):
                    return False
                if len(words) < 2 or len(words) > 4:
                    return False
                return True
                
            return False

        for table in tables:
            self.stats["tables_processed"] += 1
            
            # Map column indices to PII types based on row 0 headers
            col_mappings = {}
            if len(table.rows) > 0:
                header_row = table.rows[0]
                for idx, cell in enumerate(header_row.cells):
                    header_text = cell.text.strip().lower()
                    if "address" in header_text:
                        col_mappings[idx] = "ADDRESS"
                    elif "name" in header_text:
                        col_mappings[idx] = "FULL_NAME"
                    elif "email" in header_text:
                        col_mappings[idx] = "EMAIL"
                    elif "phone" in header_text or "telephone" in header_text:
                        col_mappings[idx] = "PHONE"
                        
            # Process rows
            for row_idx, row in enumerate(table.rows):
                # We skip checking column mapping on header row to preserve headers
                is_header = (row_idx == 0)
                processed_cells = set()
                
                for col_idx, cell in enumerate(row.cells):
                    # Deduplicate merged cells using their internal XML element tc as key
                    cell_key = cell._tc
                    if cell_key in processed_cells:
                        continue
                    processed_cells.add(cell_key)
                    
                    self.stats["cells_processed"] += 1
                    
                    # Heuristic 1: If this cell is under a mapped column header, redact its entire text
                    if not is_header and col_idx in col_mappings:
                        pii_type = col_mappings[col_idx]
                        original_text = cell.text.strip()
                        if validate_cell_text(original_text, pii_type):
                            norm_text = re.sub(r"\s+", " ", original_text)
                            if any(fv in norm_text or norm_text in fv for fv in self.redactor.generated_values):
                                continue
                            repl = self.redactor.get_replacement(original_text, pii_type)
                            # Replace all text in the cell's paragraphs
                            for p in cell.paragraphs:
                                replace_span_in_runs(p.runs, 0, len(p.text), repl)
                                
                            self.stats["replacements"] += 1
                            self.stats["detections_by_category"][pii_type] = \
                                self.stats["detections_by_category"].get(pii_type, 0) + 1
                            continue
                            
                    # Heuristic 2: Otherwise, perform regular inline regex/rule detection on cell paragraphs
                    self.process_paragraphs(cell.paragraphs)
